from utils.logger import Logger
from docx import Document
from io import BytesIO

logger = Logger.get_logger()


class TranslateDocxClient:
    def translate_paragraph(self, paragraph, translate_file_text_function):
        # Use strip to remove leading/trailing whitespace
        original_text = paragraph.text.strip()
    
        # Skip translating if the text is empty or just whitespace
        if not original_text:
            return
    
        translation = translate_file_text_function(original_text)
    
        # Clear the paragraph
        paragraph.clear()
    
        # Add only the translated text in the paragraph
        paragraph.add_run(translation)
    
    def translate_table_cell(self, cell, translate_file_text_function):
        # Use strip to remove leading/trailing whitespace
        original_text = cell.text.strip()
    
        # Skip translating if the text is empty or just whitespace
        if not original_text:
            return
    
        translation = translate_file_text_function(
            original_text)
    
        # Clear the cell and insert the translated text
        cell.text = translation
        
    def translate(self, file, translate_file_text_function):
        try:
            original_document = Document(file)
            
            # Translate the text content in paragraphs
            for paragraph in original_document.paragraphs:
                self.translate_paragraph(paragraph, translate_file_text_function)
    
            # Translate the text content in tables
            for table in original_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.translate_table_cell(cell, translate_file_text_function)
            
            
            # Save the translated document to a BytesIO object
            translated_bytes_io = BytesIO()
            original_document.save(translated_bytes_io)
            translated_bytes_io.seek(0)
            
            return translated_bytes_io
        except Exception as e:
            logger.error(f'Error in translate: {e}')
            raise
